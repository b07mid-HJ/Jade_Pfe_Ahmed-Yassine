from docx import Document
import argostranslate.translate
from argostranslate import package

# Specify the path to the directory where your models are installed
custom_model_directory = r'C:\Users\Bohmid\Desktop\prod1\translation_model\translate-fr_en-1_9.argosmodel'

# Install the model
package.install_from_path(custom_model_directory)

# Now load installed languages (including those from the custom directory)
argostranslate.translate.load_installed_languages()

def extract_text_from_docx_fr(docx_file):
    """
    Extracts and translates text from a .docx file including headings and paragraphs.
    
    Args:
    - docx_file: Path to the .docx file
    
    Returns:
    - text: A string containing extracted headings and paragraphs translated from French to English
    """
    doc = Document(docx_file)
    text = ""
    installed_languages = argostranslate.translate.get_installed_languages()

    # Find the Language objects for French and English
    french_lang = None
    english_lang = None
    for lang in installed_languages:
        if lang.code.startswith('fr'):
            french_lang = lang
        elif lang.code.startswith('en'):
            english_lang = lang
        if french_lang and english_lang:
            break
    
    if not french_lang or not english_lang:
        raise Exception("Required language models (French or English) not found.")
    
    # Find the translation model from French to English
    translation_model = french_lang.get_translation(english_lang)
    
    # Translate and accumulate text from paragraphs that are not hyperlinks, tables, or figures
    for paragraph in doc.paragraphs:
        if paragraph.hyperlinks:
            continue
        if not (paragraph.text.startswith('Table') or paragraph.text.startswith('Figure')):
            translated_text = translation_model.translate(paragraph.text)
            text += f"{translated_text}\n"
    
    return text

def extract_text_from_docx(docx_file):
    """
    Extracts text from a .docx file including headings and paragraphs.
    
    Args:
    - docx_file: Path to the .docx file
    
    Returns:
    - text: A string containing extracted headings and paragraphs
    """
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        if paragraph.hyperlinks:
            continue
        if not (paragraph.text.startswith('Table') or paragraph.text.startswith('Figure')):
            text += f"{paragraph.text}\n"
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            text += f"{paragraph.text}\n"
    return text