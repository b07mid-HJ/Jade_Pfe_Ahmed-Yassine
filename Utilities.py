from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from eparse.core import get_df_from_file
import numpy as np
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import re
from docx.exceptions import InvalidSpanError

def extract_data(w,sheet_names):
    dataset = dict()
    tables = [table for table in get_df_from_file(w , sheet = sheet_names)]
            
    for tab in tables:
        tab_value = tab[0]
        tab_sheet = tab[3]  
        arr = np.array(tab_value)
        tab_title = arr[-1, 0]
        row_table = {tab_title: tab_value}
        
        if tab_sheet not in dataset:
            dataset[tab_sheet] = [row_table]
        else :
            dataset[tab_sheet].append(row_table)
    return dataset


def dataframe_to_list_without_headers_and_last_row(data):
    data_without_last_row = data.iloc[:-1]
    data_rows = data_without_last_row.values.tolist()
    cleaned_data_rows = [[str(cell) if pd.notnull(cell) else '' for cell in row] for row in data_rows]
    return cleaned_data_rows

def insert_title_docx(doc , title):
    for _ in range(10):  # Adjust number of line breaks as needed to center the title vertically
        doc.add_paragraph()

def insert_title_docx(doc , title):
            for _ in range(10):  # Adjust number of line breaks as needed to center the title vertically
                doc.add_paragraph()
            title_para = doc.add_paragraph(title)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title_para.runs[0]
            run.bold = True
            run.font.size = Pt(24)  

            doc.add_page_break()

            title_para = doc.add_paragraph(title)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title_para.runs[0]
            run.bold = True
            run.font.size = Pt(24) 

            doc.add_page_break()

def add_paragraph_smartly(doc, paragraph):
    parts = paragraph.split('**')

    p = doc.add_paragraph()
    for i, part in enumerate(parts):
        run = p.add_run(part)
        if i % 2 == 1:  # Apply bold formatting to the parts that were between **
            run.bold = True



def is_number(text):

    normalized_text = re.sub(r'\s+', '', text.replace(',', '.')).strip()
    normalized_text = re.sub(r'[ ,]', '', text.replace('%', '').strip())

    try:
        float(normalized_text)  # Attempt to convert the cleaned text to a float
        return True
    except ValueError:
        return False




def merge_vertically(table, col_idx):
    prev_cell = None
    start_row = 0
    initial_text = None  # Store the text of the first cell in a merge sequence

    for row_idx, row in enumerate(table.rows):
        cell = row.cells[col_idx]
        cell_text = cell.text.strip()
        is_numeric = is_number(cell_text)

        if not is_numeric and cell_text not in ['', '-']:
            if prev_cell is not None and cell_text == prev_cell.text:
                if row_idx == len(table.rows) - 1:
                    try:
                        table.cell(start_row, col_idx).merge(table.cell(row_idx, col_idx))
                        table.cell(start_row, col_idx).text = initial_text
                    except InvalidSpanError as e:
                        print(f'error in the table {table}')
            else:
                if prev_cell is not None and row_idx - start_row > 1:
                    try:
                        table.cell(start_row, col_idx).merge(table.cell(row_idx - 1, col_idx))
                        table.cell(start_row, col_idx).text = initial_text
                    except InvalidSpanError as e:
                        print(f'error in the table {table}')
                start_row = row_idx
                initial_text = cell_text
                prev_cell = cell
        else:
            if prev_cell is not None and row_idx - start_row > 1:
                try:
                    table.cell(start_row, col_idx).merge(table.cell(row_idx - 1, col_idx))
                    table.cell(start_row, col_idx).text = initial_text
                except InvalidSpanError as e:
                    print(f'error in the table {table}')        
            start_row = row_idx + 1
            prev_cell = None
            initial_text = None

        if row_idx == len(table.rows) - 1 and prev_cell is not None and row_idx - start_row >= 0:
            try:
                table.cell(start_row, col_idx).merge(table.cell(row_idx, col_idx))
                table.cell(start_row, col_idx).text = initial_text
            except InvalidSpanError as e :
                print(f'error in the table {table}')



def merge_horizontally(table, row_idx):
    prev_cell = None
    start_col = 0
    initial_text = None  # This will store the text of the first cell in a merge sequence
    row = table.rows[row_idx]

    for col_idx, cell in enumerate(row.cells):
        cell_text = cell.text.strip()
        is_numeric = is_number(cell_text)

        # Only proceed if the cell is not numeric and not in the ignored list
        if not is_numeric and cell_text not in ['', '-']:
            if prev_cell is not None and cell_text == prev_cell.text:
                # Continue the merge sequence
                if col_idx == len(row.cells) - 1:
                    try:
                        row.cells[start_col].merge(row.cells[col_idx])
                        row.cells[start_col].text = initial_text  # Set the text from the first cell of the sequence
                    except InvalidSpanError as e :
                        #handle an error
                        print(f'error in the table {table}')
            else:
                # New potential merge sequence detected or a single cell
                if prev_cell is not None and col_idx - start_col > 1:
                    try:
                        row.cells[start_col].merge(row.cells[col_idx - 1])
                        row.cells[start_col].text = initial_text  # Set the text from the first cell of the sequence
                    except InvalidSpanError as e :
                        #handle an error
                        print(f'error in the table {table}')                
                start_col = col_idx  # Start a new merge block
                initial_text = cell_text  # Update the text to the current cell's text as it starts a new block
                prev_cell = cell  # Update prev_cell to current cell for next iteration
        else:
            # Reset start_col and prev_cell when a cell is ignored or numeric
            if prev_cell is not None and col_idx - start_col > 1:
                try:
                    row.cells[start_col].merge(row.cells[col_idx - 1])
                    row.cells[start_col].text = initial_text  # Set the text from the first cell of the sequence
                except InvalidSpanError as e :
                    #handle an error
                    print(f'error in the table {table}')
            start_col = col_idx + 1
            prev_cell = None
            initial_text = None  # Clear initial_text as this segment should not be merged

        # Final merging for the last sequence of similar cells at the end of the row
        if col_idx == len(row.cells) - 1 and prev_cell is not None and col_idx - start_col >= 0:
            try:
                row.cells[start_col].merge(row.cells[col_idx])
                row.cells[start_col].text = initial_text  # Set the text for the merged cells
            except InvalidSpanError as e :
                #handle an error
                print(f'error in the table {table}')
