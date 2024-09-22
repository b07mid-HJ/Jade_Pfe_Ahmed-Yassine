import docx
import streamlit as st
from streamlit_sortables import sort_items
from DescGen import descriptionGenerator
from docx.oxml.shared import qn  # feel free to move these out
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from streamlit_extras.stateful_button import button 
import pandas as pd
from io import BytesIO
import Utilities as ut


"""
    File Uploader

    It's hard to test the ability to upload files in an automated way, so here you
    should test it by hand. Please upload a CSV file and make sure a table shows up
    below with its contents.
"""
def full_generation_docx():
    st.title("Welcome to Jade Private Document Generation")
    st.markdown("""---""")
    m = st.markdown("""
    <style>
    div.stButton > button:first-child {
        width: 100%;
    }
    </style>""", unsafe_allow_html=True)
    
    css = """
    <style>
    div.stButton > button:hover {
        color: rgb(31, 150, 146) !important;
        border-color: rgb(31, 150, 146) !important;
    }
    div.stButton > button:active {
    color: rgb(31, 150, 146) !important;
    border-color: rgb(31, 150, 146) !important;
    background-color: rgb(31, 150, 146) !important;
}
    div.stButton > button:focus {
    color: rgb(31, 150, 146) !important;
    border-color: rgb(31, 150, 146) !important;
}
    </style>
    """

    st.markdown(css, unsafe_allow_html=True)
    v = "Upload the Excel file"
    w = st.file_uploader("Upload the Excel file", type="xlsx")
    if w is not None:
        project = st.text_input("Enter the name of the project(required)")
        context = st.text_area("Enter the context of the document(required)", height=25, key="context")
        if context is not None :

            #from eparse.core import get_df_from_file
            from openpyxl import load_workbook
            import io
            #import numpy as np

            wb = load_workbook(w)
            sheet_names = wb.sheetnames

            
            dataset = ut.extract_data(w,sheet_names)
            original_items = []
            for sheet, tables in dataset.items():
                items = [key for table in tables for key in table.keys()]
                original_items.append({'header': sheet, 'items': items})

            sorted_items = sort_items(original_items, multi_containers=True, direction='vertical')
            

            #number of the tables
            l = len([t for k in sorted_items for t in k['items']])

        if button('Generate Document' , key = "button100"):
            # Initialize sorted_dataset to store the re-ordered dataset
            sorted_dataset = {}

            # Loop through each sorted item (sheet) in sorted_items
            mytitles = []
            for sorted_sheet in sorted_items:

                sheet_name = sorted_sheet['header']
                sorted_tables_titles = sorted_sheet['items']
                
                unsorted_tables = dataset[sheet_name]
                
                # Create a temporary dictionary to easily access tables by their title
                title_to_table_dict = {list(table.keys())[0]: table for table in unsorted_tables}
                
                # Initialize a list to hold the sorted tables for this sheet
                sorted_tables_for_sheet = []
                
                # Loop through each sorted table title and append the corresponding table
                # from title_to_table_dict to the sorted list
                for title in sorted_tables_titles:
                    sorted_table = title_to_table_dict[title]
                    sorted_tables_for_sheet.append(sorted_table)
                    mytitles.append(title)

                    sorted_dataset[sheet_name] = sorted_tables_for_sheet
                if "second_order" not in st.session_state:
                    st.session_state["second_order"] = mytitles # mathematical Application use for remain the order of the description in the list 
                # #A >> B Teta 
                teta = []
                for i , title in enumerate(mytitles):
                    teta.append(st.session_state.second_order.index(title))
                
            if "para" not in st.session_state:
                st.session_state.para = [['Hello ... !, just click the button or write something']]*l

            if "BaseTableConversation" not in st.session_state:
                st.session_state.BaseTableConversation = descriptionGenerator(context, project)
            count = 0
            for sheet, tables in sorted_dataset.items():
                st.write(f"<span style='font-size: 20px; font-weight: bold;'>{sheet}</span>", unsafe_allow_html=True)
                
                for table in tables:
                    for title, data in table.items():
                        
                        
                        if st.checkbox(title , key = f"check{count}") :

                            # Check if data is a pandas DataFrame and convert if necessary
                            if isinstance(data, pd.DataFrame):
                                data_list = ut.dataframe_to_list_without_headers_and_last_row(data)
                            else:
                                # Assuming data is already a list of lists. Exclude the last row.
                                data_list = data[:-1] if data else data

                            if data_list:
                                st.table(data_list )
                                #st.dataframe(data_list, hide_index=True)
                                text_input = st.text_area("Hello Table ! (let's begin with a verb)", height=25, key=count) 
                                
                                if st.button('Print Description', key = f"button{count+1}") :
                                    print(st.session_state.para)
                                    history_tab = st.session_state.para[teta[count]][1:]

                                    res = st.session_state.BaseTableConversation.predict(title,data_list, text_input , history_tab)
                                    res_tab = st.session_state.para[teta[count]].copy()
                                    res_tab.append(res)
                                    st.session_state.para[teta[count]] = res_tab
                                j = 0
                                for i in st.session_state.para[teta[count]]:
                                    col1, col2 = st.columns([0.9, 0.1])
                                    j+=1
                                    with col1 :
                                        message = st.chat_message("assistant")
                                        message.write(i)  # Add paragraph after each table
                                    with col2 :
                    
                                        if st.button('🗑️' , key = f"delete{count}+{j}"):
                                            st.session_state.para[teta[count]].remove(i)
                                            st.rerun()
                                
                        count+=1

            if button('Load Document' , key = 'button200'):
                    doc = docx.Document()
                    if len(project)>0 :
                        ut.insert_title_docx(doc , project)

                    # Iterate through the sorted_dataset to add data to the Word document
                    count = 0
                    for sheet, tables in sorted_dataset.items():
                        doc.add_heading(sheet, level=1) 
                        
                        for table in tables:
                            for title, data in table.items():

                                #doc.add_heading(title, level=2) 
                                caption_paragraph = doc.add_paragraph()
                                run = caption_paragraph.add_run(f'Table {count}: {title}')
                                run.bold = True 
                                caption_paragraph.style = 'Caption'

                                run = caption_paragraph.add_run()
                                fldSimple = OxmlElement('w:fldSimple')
                                fldSimple.set(qn('w:instr'), 'TOC \\h \\z \\c "Table"')
                                run._r.append(fldSimple)
                                                        
                                if isinstance(data, pd.DataFrame):
                                    data_list = ut.dataframe_to_list_without_headers_and_last_row(data)
                                else:
                                    # Assuming data is already a list of lists. Exclude the last row.
                                    # Make sure to handle NaN values and other transformations as needed.
                                    data_list = data[:-1] if data else data

                                if data_list:
                                    word_table = doc.add_table(rows=len(data_list), cols=len(data_list[0]) if data_list else 0)
                                    word_table.style = 'TableGrid'
                                    word_table.autofit = False 
                                    word_table.allow_autofit = False
                                    # Populate the table with data, excluding the last row
                                    
                                    rows = []
                                    for i, row in enumerate(data_list):
                                        rows.append(row)
                                        cells = [c for c in row if c.strip()  not in ['' ,'-',' -']]
                                        brk = len(cells) == 1
                                        for j, cell in enumerate(row):
                                            if cell.strip() in ['' ,'-']:
                                                word_table.cell(i, j).text = " "
                                            else:
                                                word_table.cell(i, j).text = cell
                                            
                                            if i == 0 or brk :
                                                clr = "#D9D9D9" if brk else "#2F5496" #header color
                                                cell_properties = word_table.cell(i, j)._element.tcPr
                                                try:
                                                    cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
                                                except IndexError:
                                                    cell_shading = OxmlElement('w:shd') 
                                                    cell_shading.set(qn('w:fill'), clr)  
                                                cell_properties.append(cell_shading)  
                                                if i == 0:
                                                    word_table.cell(i, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Blue color
                                        ut.merge_horizontally(word_table, i)
                                    for i ,col in enumerate(word_table.columns):
                                        ut.merge_vertically(word_table, i)
        
                                doc.add_paragraph()
                                if len(st.session_state.para[teta[count]]) > 1:
                                    ut.add_paragraph_smartly(doc , st.session_state.para[teta[count]][-1])  # Add paragraph after each table    
                                else :
                                    ut.add_paragraph_smartly(doc ,st.session_state.BaseTableConversation.predict(title,data_list, "", [""]))
                                doc.add_paragraph()
                                count+=1
                    
                    # Save the document
                    doc_io = BytesIO()
                    #doc_path = "mydocument.docx"  # Adjust path as necessary
                    st.write("Document generated successfully 🎉")
                    doc.save(doc_io)
                    st.download_button(
                        label="Download Document",
                        data=doc_io.getvalue(), 
                        file_name=f"{project}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )